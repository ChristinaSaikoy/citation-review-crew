from __future__ import annotations

import re
from pathlib import Path

from docx import Document

_CITE_RE = re.compile(r"\[(\d+(?:[,，\-~]\d+)*)\]")
_HEADING_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Title"}


def read_docx(file_path: str | Path) -> str:
    doc = Document(str(file_path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


def extract_cited_passages(file_path: str | Path) -> str:
    doc = Document(str(file_path))

    chapters: list[tuple[str, list[str]]] = []
    current_heading = "Preface"
    current_paras: list[str] = []
    ref_section: list[str] = []
    in_refs = False

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name if p.style else ""

        if "toc" in style.lower():
            continue

        is_heading = style in _HEADING_STYLES
        if is_heading:
            if text == "参考文献" or text.startswith("参考文献") or text.lower().startswith("reference"):
                if current_paras:
                    chapters.append((current_heading, current_paras))
                in_refs = True
                continue
            if in_refs:
                break
            if current_paras:
                chapters.append((current_heading, current_paras))
            current_heading = text
            current_paras = []
            continue

        if in_refs:
            if style.lower() == "bibliography" or _CITE_RE.match(text):
                ref_section.append(f"[{len(ref_section) + 1}] {text}")
            continue

        if _CITE_RE.search(text):
            current_paras.append(text)

    if current_paras and not in_refs:
        chapters.append((current_heading, current_paras))

    parts: list[str] = []
    total_cited = 0
    for heading, paras in chapters:
        if not paras:
            continue
        total_cited += len(paras)
        block = f"## {heading}\n\n" + "\n\n".join(paras)
        parts.append(block)

    body = "\n\n---\n\n".join(parts)

    refs_text = "\n".join(ref_section) if ref_section else "(no reference list extracted)"

    result = (
        f"# Cited passages by chapter ({total_cited} paragraphs extracted)\n\n"
        f"{body}\n\n"
        f"---\n\n"
        f"# Reference list ({len(ref_section)} entries)\n\n"
        f"{refs_text}"
    )
    return result
